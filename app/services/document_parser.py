"""
document_parser.py — Extract, chunk, and summarise uploaded documents
for the AI chat feature.

Uses pymupdf4llm for PDFs and python-docx for DOCX files.
Text is chunked with a simple recursive splitter so that large
contracts fit within the AI model's context window.
"""

import os
import logging
from typing import Optional

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md", ".csv", ".html", ".htm"}

# Chunking parameters
CHUNK_SIZE = 2000
CHUNK_OVERLAP = 200
MAX_CONTEXT_LENGTH = 50000  # max chars sent to the AI


def extract_text(file_path: str) -> str:
    """Dispatch to the correct extractor based on file extension."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type: '{ext}'. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return _extract_docx(file_path)
    elif ext in (".html", ".htm"):
        return _extract_html(file_path)
    elif ext == ".csv":
        return _extract_csv(file_path)
    else:
        return _extract_txt(file_path)


def _extract_pdf(file_path: str) -> str:
    import pymupdf4llm
    return pymupdf4llm.to_markdown(file_path)


def _extract_docx(file_path: str) -> str:
    from docx import Document

    doc = Document(file_path)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name.lower()
        if "heading 1" in style:
            lines.append(f"# {text}")
        elif "heading 2" in style:
            lines.append(f"## {text}")
        elif "heading 3" in style:
            lines.append(f"### {text}")
        else:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines)


def _extract_txt(file_path: str) -> str:
    with open(file_path, encoding="utf-8", errors="replace") as f:
        return f.read()


def _extract_html(file_path: str) -> str:
    from html.parser import HTMLParser

    class _TextExtractor(HTMLParser):
        def __init__(self):
            super().__init__()
            self._parts: list[str] = []
            self._skip = False

        def handle_starttag(self, tag, attrs):
            if tag in ("script", "style", "nav", "header", "footer"):
                self._skip = True

        def handle_endtag(self, tag):
            if tag in ("script", "style", "nav", "header", "footer"):
                self._skip = False

        def handle_data(self, data):
            if not self._skip:
                text = data.strip()
                if text:
                    self._parts.append(text)

    with open(file_path, encoding="utf-8", errors="replace") as f:
        content = f.read()

    parser = _TextExtractor()
    parser.feed(content)
    return "\n".join(parser._parts)


def _extract_csv(file_path: str) -> str:
    import csv

    lines = []
    with open(file_path, encoding="utf-8", errors="replace", newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            if not any(row.values()):
                continue
            line = " | ".join(
                f"{k.strip()}: {v.strip()}"
                for k, v in row.items()
                if k and v and v.strip()
            )
            if line:
                lines.append(line)
    return "\n".join(lines)


def chunk_text(text: str) -> list[str]:
    """Split text into overlapping chunks using recursive separators."""
    if len(text) <= CHUNK_SIZE:
        return [text]

    separators = ["\n\n", "\n", ". ", " "]
    return _recursive_split(text, separators)


def _recursive_split(text: str, separators: list[str]) -> list[str]:
    """Split text by the first separator that produces chunks, then recurse."""
    if len(text) <= CHUNK_SIZE:
        return [text]

    chunks = []
    separator = separators[0] if separators else ""
    remaining_separators = separators[1:] if len(separators) > 1 else [""]

    parts = text.split(separator) if separator else [text[i:i + CHUNK_SIZE] for i in range(0, len(text), CHUNK_SIZE)]

    current = ""
    for part in parts:
        candidate = f"{current}{separator}{part}" if current else part
        if len(candidate) <= CHUNK_SIZE:
            current = candidate
        else:
            if current:
                if len(current) > CHUNK_SIZE:
                    chunks.extend(_recursive_split(current, remaining_separators))
                else:
                    chunks.append(current)
            current = part

    if current:
        if len(current) > CHUNK_SIZE:
            chunks.extend(_recursive_split(current, remaining_separators))
        else:
            chunks.append(current)

    # Add overlap between chunks
    if CHUNK_OVERLAP > 0 and len(chunks) > 1:
        overlapped = [chunks[0]]
        for i in range(1, len(chunks)):
            prev_tail = chunks[i - 1][-CHUNK_OVERLAP:]
            overlapped.append(prev_tail + chunks[i])
        chunks = overlapped

    return chunks


def build_contract_context(text: str) -> str:
    """
    Build a condensed context from the extracted document text.
    Chunks the text and takes the most important parts to stay
    within the AI's context limits.
    """
    if not text or not text.strip():
        return ""

    # If short enough, use the full text
    if len(text) <= MAX_CONTEXT_LENGTH:
        return text

    # Chunk and take the beginning and end (contract preamble + signature/terms)
    chunks = chunk_text(text)
    if not chunks:
        return ""

    context_parts = []
    current_length = 0

    # Take chunks from the beginning (usually has parties, dates, definitions)
    for chunk in chunks:
        if current_length + len(chunk) > MAX_CONTEXT_LENGTH:
            break
        context_parts.append(chunk)
        current_length += len(chunk)

    return "\n\n".join(context_parts)


def extract_key_info_prompt(document_text: str) -> str:
    """
    Build a prompt section that asks the AI to first identify key
    contract information before answering the user's question.
    """
    return (
        "The following is an uploaded contract document. "
        "First, identify the key information in this contract including: "
        "parties involved, contract type, effective dates, key terms, "
        "payment terms, obligations, termination conditions, and any "
        "notable clauses or risks. Then use this understanding to answer "
        "the user's question.\n\n"
        "--- DOCUMENT START ---\n"
        f"{document_text}\n"
        "--- DOCUMENT END ---"
    )
