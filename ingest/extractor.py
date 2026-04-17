"""
extractor.py — Universal text extractor.

Converts any supported file format into a single plain-text string
that can then be chunked and embedded by the ingestion pipeline.

Supported formats:
    .txt              — read as-is
    .pdf              — pymupdf4llm -> Markdown
    .csv              — rows joined as "col: value, col: value" lines
    .json             — flattened to readable key: value text
    .docx             — python-docx paragraph extraction
    .xlsx             — openpyxl sheet -> CSV-style text
    .html / .htm      — BeautifulSoup text extraction
    .md               — read as-is (Markdown is plain text)

Usage:
    from extractor import extract
    text = extract("/data/input/contracts.txt")
"""

import os
import json
import csv
import io
import logging

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {
    ".txt", ".pdf", ".csv", ".json",
    ".docx", ".xlsx", ".html", ".htm", ".md"
}


def extract(file_path: str) -> str:
    """
    Dispatch to the correct extractor based on file extension.
    Returns a single plain-text string.
    Raises ValueError for unsupported extensions.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type: '{ext}'. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    extractors = {
        ".txt":  _extract_txt,
        ".md":   _extract_txt,       # Markdown is plain text
        ".pdf":  _extract_pdf,
        ".csv":  _extract_csv,
        ".json": _extract_json,
        ".docx": _extract_docx,
        ".xlsx": _extract_xlsx,
        ".html": _extract_html,
        ".htm":  _extract_html,
    }

    logger.info("Format: %s", ext.lstrip('.'))
    return extractors[ext](file_path)


# -- Individual extractors --------------------------------------------------


def _extract_txt(file_path: str) -> str:
    """Plain text and Markdown -- read directly."""
    with open(file_path, encoding="utf-8", errors="replace") as f:
        return f.read()


def _extract_pdf(file_path: str) -> str:
    """PDF -- convert to Markdown via pymupdf4llm (preserves structure)."""
    try:
        import pymupdf4llm
    except ImportError:
        raise ImportError("pymupdf4llm is required for PDF extraction. Run: pip install pymupdf4llm")
    return pymupdf4llm.to_markdown(file_path)


def _extract_csv(file_path: str) -> str:
    """
    CSV -- convert each row to a human-readable sentence.

    Input row:  customer, value, date
                Acme,     50000, 2024-01-01

    Output:     customer: Acme | value: 50000 | date: 2024-01-01
    """
    lines = []
    with open(file_path, encoding="utf-8", errors="replace", newline="") as f:
        reader = csv.DictReader(f)
        for i, row in enumerate(reader):
            # Skip completely empty rows
            if not any(row.values()):
                continue
            line = " | ".join(
                f"{k.strip()}: {v.strip()}"
                for k, v in row.items()
                if k and v and v.strip()
            )
            if line:
                lines.append(line)
    return "\n".join(lines)


def _extract_json(file_path: str) -> str:
    """
    JSON -- flatten nested structure into readable text.

    Handles:
      - Array of objects  -> one paragraph per object
      - Single object     -> key: value pairs
      - Array of strings  -> joined lines
      - Nested objects    -> recursively flattened with dot-notation keys
    """
    with open(file_path, encoding="utf-8", errors="replace") as f:
        data = json.load(f)

    def flatten(obj, prefix="") -> list[str]:
        parts = []
        if isinstance(obj, dict):
            for k, v in obj.items():
                key = f"{prefix}.{k}" if prefix else k
                parts.extend(flatten(v, key))
        elif isinstance(obj, list):
            for i, v in enumerate(obj):
                parts.extend(flatten(v, f"{prefix}[{i}]" if prefix else f"[{i}]"))
        else:
            val = str(obj).strip()
            if val:
                parts.append(f"{prefix}: {val}" if prefix else val)
        return parts

    if isinstance(data, list):
        blocks = []
        for item in data:
            if isinstance(item, dict):
                block = "\n".join(flatten(item))
            elif isinstance(item, str):
                block = item.strip()
            else:
                block = str(item)
            if block:
                blocks.append(block)
        return "\n\n".join(blocks)

    return "\n".join(flatten(data))


def _extract_docx(file_path: str) -> str:
    """
    DOCX -- extract paragraphs and table cells via python-docx.
    Preserves heading structure with Markdown-style prefixes.
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        raise ImportError("python-docx is required for DOCX extraction. Run: pip install python-docx")

    doc = Document(file_path)
    lines = []

    # Paragraphs (headings get # prefix for structure)
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name.lower()
        if "heading 1" in style:
            lines.append(f"# {text}")
        elif "heading 2" in style:
            lines.append(f"## {text}")
        elif "heading 3" in style:
            lines.append(f"### {text}")
        else:
            lines.append(text)

    # Tables -- each row as pipe-separated values
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(lines)


def _extract_xlsx(file_path: str) -> str:
    """
    XLSX -- extract all sheets as readable text.
    Uses the first row of each sheet as column headers.
    Each subsequent row becomes a "header: value | ..." line.
    """
    try:
        import openpyxl
    except ImportError:
        raise ImportError("openpyxl is required for XLSX extraction. Run: pip install openpyxl")

    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    blocks = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        # First row as headers
        headers = [str(h).strip() if h is not None else f"col_{i}"
                   for i, h in enumerate(rows[0])]

        sheet_lines = [f"Sheet: {sheet_name}"]
        for row in rows[1:]:
            if all(v is None for v in row):
                continue
            line = " | ".join(
                f"{headers[i]}: {str(v).strip()}"
                for i, v in enumerate(row)
                if v is not None and str(v).strip()
            )
            if line:
                sheet_lines.append(line)

        blocks.append("\n".join(sheet_lines))

    wb.close()
    return "\n\n".join(blocks)


def _extract_html(file_path: str) -> str:
    """
    HTML -- strip all tags with BeautifulSoup, keeping visible text only.
    Removes script, style, nav, header, footer noise.
    """
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise ImportError("beautifulsoup4 is required for HTML extraction. Run: pip install beautifulsoup4")

    with open(file_path, encoding="utf-8", errors="replace") as f:
        soup = BeautifulSoup(f.read(), "html.parser")

    # Remove non-content tags
    for tag in soup(["script", "style", "nav", "header", "footer",
                     "aside", "noscript", "meta", "link"]):
        tag.decompose()

    # Get text, collapsing excess whitespace
    text = soup.get_text(separator="\n")
    lines = [line.strip() for line in text.splitlines()]
    lines = [line for line in lines if line]          # drop blank lines
    return "\n".join(lines)
